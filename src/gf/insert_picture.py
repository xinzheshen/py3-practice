# 引入库
import os
import re
from argparse import ArgumentParser

from docx import Document
from docx.shared import Cm
import xml.etree.cElementTree as ET


file_path = r"D:\work\gf\截图\15.第十五 心268-test.docx"
pic_file_path = r"D:\work\gf\15-元心.png"

# # 打开文档1
# doc = Document(file_path)
#
# # 读取每段内容
# pl = [paragraph.text for paragraph in doc.paragraphs]
# pre_p = None
# for i, p in enumerate(doc.paragraphs):  # 遍历所有的段落
#         print(str(i) + ":"+ str(p.text))
#         if len(p.text) != 0:
#
#             for i in range(len(p.runs)):  # p.runs代表p这个段落下所有文字的列表
#                 print(str(i)+':::::')
#                 print(p.runs[i].text)  # 当打印时，发现p.runs把段落自动分解了
#         if str(1) in p.text:
#             pre_p.runs[-1].add_break()  # 添加一个折行
#             inline_shape = pre_p.runs[-1].add_picture(pic_file_path, height=2400000)  # 在runs的最后一段文字后添加图片
#             # inline_shape.height = Cm(7)
#
#             # p.runs[-1].add_break()  # 添加一个折行
#             # p.runs[-1].add_picture(pic_file_path)  # 在runs的最后一段文字后添加图片
#             # os.remove(photo_dit_path)
#             doc.save(r"D:\work\gf\15.第十五 心268-test123.docx")  # 保存文件
#             break
#
#         pre_p = p


def load_picture_list(pic_dir):
    pic_dict = {}
    for file in os.listdir(pic_dir):
        file_key = file.split("_")[-1].split(".")[0]
        file_path = os.path.join(pic_dir, file)
        pic_dict[file_key] = file_path
    return pic_dict


def insert_picture(word_path, pic_dict):
    pattern = re.compile(r'\d+\.')

    tmp = word_path.rfind(".")
    word_file_save_path = word_path[:tmp] + "_new" + word_path[tmp:]

    # 打开文档1
    doc = Document(word_path)

    pre_p = None
    for i, p in enumerate(doc.paragraphs):  # 遍历所有的段落
        if pattern.match(p.text):
            # pic_key = p.text[:-1]
            pic_key = p.text.split(".")[0]
            pic_file_path = pic_dict.get(pic_key)
            if pic_file_path is None:
                # todo 记录
                print(f"没有第{pic_key}张图片。")
                continue

            pre_p.runs[-1].add_break()  # 添加一个折行
            inline_shape = pre_p.runs[-1].add_picture(pic_file_path, height=picture_height)  # 在runs的最后一段文字后添加图片
            # inline_shape.height = Cm(7)

        pre_p = p

    doc.save(word_file_save_path)  # 保存文件


def process_argv():
    parser = ArgumentParser(prog='insert_picture')
    parser.add_argument('--word_file', help='PDF文件路径', type=str, required=True)
    parser.add_argument('--picture_path', help='截图存放路径', type=str, required=True)
    parser.add_argument('--picture_height', help='插入图片高度', type=float, default=2400000, required=False)

    return parser.parse_args()


if __name__ == '__main__':
    arg = process_argv()

    word_file_path = arg.word_file
    picture_dir = arg.picture_path
    picture_height = arg.picture_height

    # word_file_path = r"D:\work\gf\截图\16.第十六 忄篇540.docx"
    # picture_dir = r"D:\work\gf\截图\pictures\16"
    # picture_height = 2400000

    picture_dict = load_picture_list(picture_dir)

    insert_picture(word_file_path, picture_dict)
